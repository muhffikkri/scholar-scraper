"""
File handler module for Google Scholar scraper.
Menangani operasi baca/tulis file untuk berbagai format (CSV, TXT, XLSX, DOCX).
"""

import csv
import os
from typing import List, Dict, Optional
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# Load environment variables dari file .env
load_dotenv()


def get_config(key: str, default: Optional[str] = None) -> Optional[str]:
    """
    Mendapatkan nilai konfigurasi dari environment variable (.env file).
    
    Args:
        key (str): Nama konfigurasi yang ingin diambil
        default (Optional[str]): Nilai default jika konfigurasi tidak ditemukan
        
    Returns:
        Optional[str]: Nilai konfigurasi atau default
    """
    return os.getenv(key, default)


def read_dosen_from_file(filepath: str) -> List[str]:
    """
    Membaca daftar nama dosen dari file CSV atau TXT.
    
    Args:
        filepath (str): Path ke file yang berisi daftar nama dosen
        
    Returns:
        List[str]: List berisi nama-nama dosen
        
    Raises:
        FileNotFoundError: Jika file tidak ditemukan
        ValueError: Jika format file tidak didukung
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File tidak ditemukan: {filepath}")
    
    file_extension = os.path.splitext(filepath)[1].lower()
    dosen_names = []
    
    try:
        if file_extension == '.csv':
            with open(filepath, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                # Skip header jika ada
                header = next(reader, None)
                for row in reader:
                    if row and row[0].strip():
                        dosen_names.append(row[0].strip())
        
        elif file_extension == '.txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line:
                        dosen_names.append(line)
        
        else:
            raise ValueError(f"Format file tidak didukung: {file_extension}. Gunakan .csv atau .txt")
    
    except UnicodeDecodeError:
        # Coba dengan encoding yang berbeda
        try:
            if file_extension == '.csv':
                with open(filepath, 'r', encoding='latin-1') as f:
                    reader = csv.reader(f)
                    header = next(reader, None)
                    for row in reader:
                        if row and row[0].strip():
                            dosen_names.append(row[0].strip())
            elif file_extension == '.txt':
                with open(filepath, 'r', encoding='latin-1') as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            dosen_names.append(line)
        except Exception as e:
            raise ValueError(f"Gagal membaca file dengan berbagai encoding: {e}")
    
    return dosen_names


def save_to_csv(df: pd.DataFrame, filename: str) -> str:
    """
    Menyimpan DataFrame ke file CSV.
    
    Args:
        df (pd.DataFrame): DataFrame berisi data publikasi
        filename (str): Nama file output (dengan atau tanpa ekstensi .csv)
        
    Returns:
        str: Path lengkap file yang disimpan
    """
    if not filename.endswith('.csv'):
        filename += '.csv'
    
    # Pastikan direktori output ada
    output_dir = os.path.dirname(filename)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    df.to_csv(filename, index=False, encoding='utf-8-sig')
    return os.path.abspath(filename)


def save_to_excel(df: pd.DataFrame, filename: str) -> str:
    """
    Menyimpan DataFrame ke file Excel dengan formatting.
    
    Args:
        df (pd.DataFrame): DataFrame berisi data publikasi
        filename (str): Nama file output (dengan atau tanpa ekstensi .xlsx)
        
    Returns:
        str: Path lengkap file yang disimpan
    """
    if not filename.endswith('.xlsx'):
        filename += '.xlsx'
    
    # Pastikan direktori output ada
    output_dir = os.path.dirname(filename)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Simpan dengan formatting
    with pd.ExcelWriter(filename, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Publications')
        
        # Dapatkan worksheet untuk formatting
        worksheet = writer.sheets['Publications']
        
        # Auto-adjust kolom width
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
    
    return os.path.abspath(filename)


def generate_summary_docx(df: pd.DataFrame, filename: str) -> str:
    """
    Membuat dokumen Word berisi ringkasan publikasi.
    
    Args:
        df (pd.DataFrame): DataFrame berisi data publikasi
        filename (str): Nama file output (dengan atau tanpa ekstensi .docx)
        
    Returns:
        str: Path lengkap file yang disimpan
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    # Pastikan direktori output ada
    output_dir = os.path.dirname(filename)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Buat dokumen baru
    doc = Document()
    
    # Tambahkan judul
    title = doc.add_heading('Ringkasan Publikasi Google Scholar', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tambahkan informasi umum
    doc.add_paragraph(f'Total Publikasi: {len(df)}')
    doc.add_paragraph(f'Tanggal Pembuatan: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()
    
    # Group by dosen
    if 'Nama Dosen' in df.columns:
        grouped = df.groupby('Nama Dosen')
        
        for dosen_name, group in grouped:
            # Tambahkan heading untuk setiap dosen
            doc.add_heading(f'{dosen_name}', level=1)
            doc.add_paragraph(f'Jumlah Publikasi: {len(group)}')
            doc.add_paragraph()
            
            # Tambahkan tabel publikasi
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header tabel
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'No'
            hdr_cells[1].text = 'Judul'
            hdr_cells[2].text = 'Venue'
            hdr_cells[3].text = 'Tahun'
            
            # Isi tabel
            for idx, (_, row) in enumerate(group.iterrows(), 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = str(row.get('Judul', ''))
                row_cells[2].text = str(row.get('Venue', ''))
                row_cells[3].text = str(row.get('Tahun', ''))
            
            doc.add_paragraph()
    else:
        # Jika tidak ada kolom nama dosen, buat tabel sederhana
        doc.add_heading('Semua Publikasi', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[1].text = 'Judul'
        hdr_cells[2].text = 'Venue'
        hdr_cells[3].text = 'Tahun'
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = str(row.get('Judul', ''))
            row_cells[2].text = str(row.get('Venue', ''))
            row_cells[3].text = str(row.get('Tahun', ''))
    
    # Simpan dokumen
    doc.save(filename)
    return os.path.abspath(filename)


def ensure_output_directory(directory: str = 'output') -> str:
    """
    Memastikan direktori output ada, jika tidak akan dibuat.
    
    Args:
        directory (str): Nama direktori output
        
    Returns:
        str: Path lengkap direktori output
    """
    if not os.path.exists(directory):
        os.makedirs(directory)
    return os.path.abspath(directory)


def transfer_data_to_sheets(
    excel_file_path: str,
    spreadsheet_url: str,
    sheet_name: str = None,
    web_app_url: str = None,
    status_callback=None
) -> dict:
    """
    Transfer data dari file Excel lokal ke Google Spreadsheet melalui Apps Script Web API.
    
    Args:
        excel_file_path (str): Path ke file Excel yang akan diupload
        spreadsheet_url (str): URL Google Spreadsheet tujuan
        sheet_name (str, optional): Nama sheet di Google Spreadsheet. 
                                    Jika None, akan menggunakan DEFAULT_SHEET_NAME dari .env
        web_app_url (str, optional): URL Web App dari Apps Script.
                                     Jika None, akan menggunakan APPS_SCRIPT_URL dari .env
        status_callback (callable, optional): Fungsi callback untuk update status
        
    Returns:
        dict: Response dari API dengan status dan pesan
        
    Raises:
        FileNotFoundError: Jika file Excel tidak ditemukan
        ValueError: Jika format file tidak valid atau konfigurasi tidak lengkap
    """
    import requests
    import json
    from src.core_logic.utils import extract_spreadsheet_id_from_url
    
    # Gunakan nilai dari .env jika parameter tidak diberikan
    if web_app_url is None:
        web_app_url = get_config('APPS_SCRIPT_URL')
        if not web_app_url or web_app_url == 'https://script.google.com/macros/s/YOUR_SCRIPT_ID_HERE/exec':
            raise ValueError(
                "APPS_SCRIPT_URL tidak ditemukan atau belum dikonfigurasi.\n"
                "Silakan isi APPS_SCRIPT_URL di file .env\n"
                "Gunakan .env.example sebagai template."
            )
    
    if sheet_name is None:
        sheet_name = get_config('DEFAULT_SHEET_NAME', 'Publikasi Dosen')
    
    def log(message):
        if status_callback:
            status_callback(message)
    
    # Validasi file
    if not os.path.exists(excel_file_path):
        raise FileNotFoundError(f"File tidak ditemukan: {excel_file_path}")
    
    if not excel_file_path.endswith('.xlsx'):
        raise ValueError("File harus berformat .xlsx")
    
    log("📖 Membaca data dari file Excel...")
    
    # Baca file Excel
    try:
        df = pd.read_excel(excel_file_path)
        log(f"✅ Berhasil membaca {len(df)} baris data")
    except Exception as e:
        raise ValueError(f"Gagal membaca file Excel: {e}")
    
    # Ekstrak Spreadsheet ID
    log("🔍 Mengekstrak Spreadsheet ID dari URL...")
    spreadsheet_id = extract_spreadsheet_id_from_url(spreadsheet_url)
    
    if not spreadsheet_id:
        raise ValueError("URL Google Spreadsheet tidak valid")
    
    log(f"✅ Spreadsheet ID: {spreadsheet_id}")
    
    # Konversi DataFrame ke format yang dikirim ke API
    log("📦 Menyiapkan data untuk transfer...")
    
    # Header
    headers = df.columns.tolist()
    
    # Data rows
    data_rows = df.values.tolist()
    
    # Convert semua nilai ke string untuk menghindari masalah JSON
    data_rows = [[str(cell) if pd.notna(cell) else "" for cell in row] for row in data_rows]
    
    # Gabungkan header dan data
    all_data = [headers] + data_rows
    
    log(f"📊 Total kolom: {len(headers)}")
    log(f"📊 Total baris data: {len(data_rows)}")
    
    # Siapkan payload untuk API
    payload = {
        "spreadsheetId": spreadsheet_id,
        "sheetName": sheet_name,
        "data": all_data
    }
    
    log("🚀 Mengirim data ke Google Sheets...")
    log(f"   Target: {sheet_name}")
    
    try:
        # Kirim POST request ke Apps Script Web App
        response = requests.post(
            web_app_url,
            json=payload,
            headers={"Content-Type": "application/json"},
            timeout=60
        )
        
        response.raise_for_status()
        result = response.json()
        
        if result.get("status") == "success":
            log("✅ SUKSES: Data berhasil ditulis ke Google Sheets!")
            log(f"   Spreadsheet ID: {spreadsheet_id}")
            log(f"   Sheet: {sheet_name}")
            log(f"   Baris ditulis: {len(all_data)}")
            return result
        else:
            error_msg = result.get("message", "Unknown error")
            log(f"❌ GAGAL: {error_msg}")
            raise Exception(error_msg)
            
    except requests.exceptions.Timeout:
        log("❌ ERROR: Request timeout (>60 detik)")
        raise Exception("Request timeout. Coba lagi atau periksa koneksi internet.")
    
    except requests.exceptions.RequestException as e:
        log(f"❌ ERROR: Gagal menghubungi server - {e}")
        raise Exception(f"Gagal menghubungi server: {e}")
    
    except json.JSONDecodeError:
        log("❌ ERROR: Response dari server tidak valid")
        raise Exception("Response dari server tidak valid (bukan JSON)")
